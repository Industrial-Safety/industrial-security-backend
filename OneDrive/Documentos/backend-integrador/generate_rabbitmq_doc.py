from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def para(text='', bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def code_block(text):
    """Párrafo con fuente monoespaciada y fondo gris simulado."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    # Sombreado gris
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, (text, width) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = header
            run.font.size = Pt(10)
            if header:
                run.font.color.rgb = RGBColor(255, 255, 255)
    if header:
        for cell in row.cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'),   'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'),  '1F3864')
            cell._tc.get_or_add_tcPr().append(shading)
    return row

def zebra_row(row, odd=True):
    fill = 'EBF3FB' if odd else 'FFFFFF'
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'),   'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'),  fill)
        cell._tc.get_or_add_tcPr().append(shading)

# ═════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Arquitectura de Mensajería')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(31, 56, 100)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('RabbitMQ — Industrial Safety Backend')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(70, 130, 180)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('Flujos de eventos entre microservicios')
r3.italic = True
r3.font.size = Pt(13)
r3.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. CONCEPTO BASE
# ═════════════════════════════════════════════════════════════════════════════
heading('1. Concepto base', 1, color=(31, 56, 100))

para('El patrón de mensajería del proyecto sigue la arquitectura Publish/Subscribe con RabbitMQ como broker central. '
     'Los microservicios no se llaman entre sí directamente — publican eventos y otros servicios los consumen de forma asíncrona.', size=11)

doc.add_paragraph()
para('Componentes principales:', bold=True)

items = [
    ('Producer (Publisher)', 'Microservicio que publica un evento al exchange.'),
    ('Exchange',             'Pieza de RabbitMQ (no un microservicio) que recibe el mensaje y lo enruta a colas según routing keys.'),
    ('Cola (Queue)',         'Buffer donde se almacena el mensaje hasta que el consumer lo procese.'),
    ('Consumer (Listener)',  'Microservicio que escucha una cola y ejecuta lógica cuando llega un mensaje.'),
    ('Dead Letter Queue',    'Cola de seguridad donde van los mensajes que fallaron 3 veces. Nunca se pierde un mensaje.'),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.columns[0].width = Cm(5)
tbl.columns[1].width = Cm(11.5)
table_row(tbl, [('Componente', 5), ('Descripción', 11.5)], header=True)
for i, (comp, desc) in enumerate(items):
    row = tbl.add_row()
    row.cells[0].text = comp
    row.cells[1].text = desc
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
    zebra_row(row, odd=(i % 2 == 0))

doc.add_paragraph()
code_block(
    "PRODUCER          RABBITMQ BROKER                     CONSUMER\n"
    "(publica)    →   [Exchange → routing key → Cola]  →  (escucha / @RabbitListener)"
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 2. EXCHANGES DEL PROYECTO
# ═════════════════════════════════════════════════════════════════════════════
heading('2. Exchanges del proyecto', 1, color=(31, 56, 100))

para('El proyecto tiene dos exchanges, cada uno con su propósito:', size=11)
doc.add_paragraph()

heading('2.1  industrial.safety.topic  (TopicExchange)', 2, color=(70, 130, 180))
para('Compartido por todos los microservicios principales. Usa wildcards en el routing key '
     '(# = cualquier sufijo), lo que permite que un solo mensaje pueda enrutar a múltiples colas '
     'según el patrón.', size=11)

doc.add_paragraph()
code_block(
    "EXCHANGE: industrial.safety.topic  (TopicExchange)\n"
    "│\n"
    "├─ event.order.created      ──► payment.order.created.queue   (payment-service consume)\n"
    "├─ event.payment.success    ──► order.payment.result.queue    (order-service consume)\n"
    "├─ event.payment.failed     ──► order.payment.result.queue    (order-service consume)\n"
    "├─ event.email.#            ──► notification.email.queue      (notification-service consume)\n"
    "├─ event.alert.#            ──► notification.ws.alert.queue   (notification-service consume)\n"
    "└─ event.certificate.#      ──► notification.certificate.queue (notification-service consume)"
)

doc.add_paragraph()
heading('2.2  epp.exchange  (DirectExchange)', 2, color=(70, 130, 180))
para('Exclusivo del purchase-service. Usa DirectExchange (routing key exacto, sin wildcards). '
     'Más simple porque solo tiene un tipo de evento.', size=11)

code_block(
    "EXCHANGE: epp.exchange  (DirectExchange)\n"
    "│\n"
    "└─ epp.delivered  ──► epp.delivery.queue  (sin consumer hoy — listo para notification-service)"
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3. FLUJOS COMPLETOS
# ═════════════════════════════════════════════════════════════════════════════
heading('3. Flujos de eventos por caso de uso', 1, color=(31, 56, 100))

# ── Flujo 1 ──────────────────────────────────────────────────────────────────
heading('3.1  Compra de curso (flujo más completo)', 2, color=(70, 130, 180))
para('Participan: order-service, payment-service, notification-service', italic=True, color=(100,100,100))
doc.add_paragraph()
code_block(
    "ALUMNO compra un curso\n"
    "        │\n"
    "        ▼\n"
    "[order-service]\n"
    "  1. Crea Order en DB (estado: PENDING)\n"
    "  2. Publica → 'event.order.created'\n"
    "        │\n"
    "        ▼  (RabbitMQ enruta a payment.order.created.queue)\n"
    "        │\n"
    "[payment-service]  ← @RabbitListener\n"
    "  3. Recibe OrderCreatedEvent\n"
    "  4. Crea preferencia en MercadoPago API\n"
    "  5. Alumno paga en MercadoPago (externo)\n"
    "  6. MercadoPago notifica via webhook → /api/v1/payments/webhook\n"
    "  7. Publica → 'event.payment.success' o 'event.payment.failed'\n"
    "        │\n"
    "        ▼  (RabbitMQ enruta a order.payment.result.queue)\n"
    "        │\n"
    "[order-service]  ← @RabbitListener\n"
    "  8. Recibe PaymentResultEvent\n"
    "  9. Actualiza Order en DB (estado: PAID o FAILED)\n"
    " 10. Publica → 'event.email.purchase.success/failed'\n"
    " 11. Publica → 'event.alert.purchase.success/failed'\n"
    "        │\n"
    "        ├──► notification.email.queue\n"
    "        │      [notification-service] ← @RabbitListener\n"
    "        │        Envía email al alumno (éxito o fallo)\n"
    "        │\n"
    "        └──► notification.ws.alert.queue\n"
    "               [notification-service] ← @RabbitListener\n"
    "                 Push WebSocket al alumno en tiempo real"
)

doc.add_paragraph()

# ── Flujo 2 ──────────────────────────────────────────────────────────────────
heading('3.2  Alumno aprueba examen → Certificado', 2, color=(70, 130, 180))
para('Participan: exam-service, notification-service', italic=True, color=(100,100,100))
doc.add_paragraph()
code_block(
    "[exam-service]\n"
    "  1. Alumno completa intento y aprueba el examen\n"
    "  2. Genera PDF del certificado (S3)\n"
    "  3. Guarda Certificate en DB\n"
    "  4. Publica → 'event.certificate.passed'\n"
    "        │\n"
    "        ▼  (RabbitMQ enruta a notification.certificate.queue)\n"
    "        │\n"
    "[notification-service]  ← @RabbitListener\n"
    "  5. Recibe CertificateEmailRequest\n"
    "  6. Envía email al alumno con el certificado adjunto"
)

doc.add_paragraph()

# ── Flujo 3 ──────────────────────────────────────────────────────────────────
heading('3.3  Instructor solicita cambio de precio', 2, color=(70, 130, 180))
para('Participan: course-service, notification-service', italic=True, color=(100,100,100))
doc.add_paragraph()
code_block(
    "[course-service]\n"
    "  1. Instructor crea PriceChangeRequest\n"
    "  2. Publica → 'event.email.price_request'\n"
    "        │\n"
    "        ▼  notification.email.queue\n"
    "[notification-service]\n"
    "  3. Envía email a gerencia notificando la solicitud\n"
    "\n"
    "  ── Gerencia revisa y aprueba o rechaza ──\n"
    "\n"
    "[course-service]\n"
    "  4a. Si APROBADO → Publica → 'event.alert.price_approved'\n"
    "  4b. Si RECHAZADO → Publica → 'event.alert.price_rejected'\n"
    "        │\n"
    "        ▼  notification.ws.alert.queue\n"
    "[notification-service]\n"
    "  5. Push WebSocket al instructor con el resultado"
)

doc.add_paragraph()

# ── Flujo 4 ──────────────────────────────────────────────────────────────────
heading('3.4  Entrega de EPP a trabajador', 2, color=(70, 130, 180))
para('Participan: purchase-service  (consumer pendiente de implementar)', italic=True, color=(100,100,100))
doc.add_paragraph()
code_block(
    "[purchase-service]\n"
    "  1. Logístico busca trabajador por DNI (llama a user-service via RestTemplate)\n"
    "  2. Selecciona EPP del inventario\n"
    "  3. Confirma entrega → reduce stock en DB\n"
    "  4. Guarda EppDelivery en DB\n"
    "  5. Publica → 'epp.delivered' al exchange 'epp.exchange'\n"
    "        │\n"
    "        ▼  epp.delivery.queue\n"
    "  [ Sin consumer actualmente — mensaje guardado en cola ]\n"
    "\n"
    "  ── Implementación futura ──────────────────────────────\n"
    "  notification-service podría consumir epp.delivery.queue\n"
    "  y enviar email/push al trabajador con el detalle del EPP entregado"
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 4. DEAD LETTER QUEUE
# ═════════════════════════════════════════════════════════════════════════════
heading('4. Dead Letter Queue (DLQ) — Seguro de vida', 1, color=(31, 56, 100))

para('Todas las colas principales tienen configurado un DLQ. Si el consumer falla al procesar un mensaje '
     '(lanza excepción no controlada), el mensaje se mueve a la cola DLQ después de los reintentos. '
     'Nunca se pierde un mensaje.', size=11)

doc.add_paragraph()
code_block(
    "Cola principal falla (RuntimeException no controlada)\n"
    "  channel.basicNack(tag, false, false)  ← no re-encolar\n"
    "        │\n"
    "        ▼\n"
    "[industrial.safety.dlx]  ← Dead Letter Exchange\n"
    "        │\n"
    "        ▼\n"
    "Cola DLQ (ej: notification.email.dlq)\n"
    "  Mensaje guardado aquí para revisión/reproceso manual\n"
    "\n"
    "Si el error es de dominio controlado (ej: orden ya procesada):\n"
    "  channel.basicAck(tag, false)  ← ack normal, descartar limpiamente"
)

doc.add_paragraph()

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
table_row(tbl2, [('Cola principal', 5), ('Cola DLQ', 5), ('Consumer', 6.5)], header=True)
dlq_data = [
    ('notification.email.queue',       'notification.email.dlq',       'notification-service'),
    ('notification.ws.alert.queue',    'notification.ws.alert.dlq',    'notification-service'),
    ('notification.certificate.queue', 'notification.certificate.dlq', 'notification-service'),
    ('payment.order.created.queue',    'payment.order.created.dlq',    'payment-service'),
    ('order.payment.result.queue',     'order.payment.result.dlq',     'order-service'),
]
for i, row_data in enumerate(dlq_data):
    row = tbl2.add_row()
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        for run in row.cells[j].paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.name = 'Courier New'
    zebra_row(row, odd=(i % 2 == 0))

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 5. TABLA RESUMEN
# ═════════════════════════════════════════════════════════════════════════════
heading('5. Resumen: quién publica y quién consume', 1, color=(31, 56, 100))

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
table_row(tbl3, [('Microservicio', 5), ('Publica (Producer)', 7), ('Consume (Consumer)', 4.5)], header=True)
summary = [
    ('order-service',        'event.order.created\nevent.email.purchase.*\nevent.alert.purchase.*', 'event.payment.*'),
    ('payment-service',      'event.payment.success\nevent.payment.failed',                         'event.order.created'),
    ('course-service',       'event.email.price_request\nevent.alert.price_approved/rejected',       '—'),
    ('exam-service',         'event.certificate.passed',                                            '—'),
    ('notification-service', '—',                                                                   'event.email.#\nevent.alert.#\nevent.certificate.#'),
    ('purchase-service',     'epp.delivered',                                                       '—  (pendiente)'),
]
for i, (svc, pub, cons) in enumerate(summary):
    row = tbl3.add_row()
    row.cells[0].text = svc
    row.cells[1].text = pub
    row.cells[2].text = cons
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
    zebra_row(row, odd=(i % 2 == 0))

doc.add_paragraph()

# ── Nota final ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('Nota: ')
r.bold = True
r.font.color.rgb = RGBColor(31, 56, 100)
p.add_run(
    'El notification-service es el único consumer puro del proyecto — no publica ningún evento, '
    'solo escucha colas y actúa (envío de email o push via WebSocket). '
    'Todos los demás servicios son producers, o producers y consumers a la vez (order-service y payment-service).'
)

# ── Pie de página ─────────────────────────────────────────────────────────────
doc.add_paragraph()
pie = doc.add_paragraph()
pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_pie = pie.add_run('Industrial Safety Backend  ·  Arquitectura RabbitMQ  ·  2026')
r_pie.font.size = Pt(9)
r_pie.font.color.rgb = RGBColor(150, 150, 150)
r_pie.italic = True

# ── Guardar ───────────────────────────────────────────────────────────────────
output = r'C:\Users\panc1\OneDrive\Documentos\backend-integrador\RabbitMQ_Flujos_IndustrialSafety.docx'
doc.save(output)
print(f'Documento generado: {output}')
