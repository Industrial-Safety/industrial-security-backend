"""
Genera la Matriz de Casos de Prueba — SafeIndustrial Backend Completo
Estructura idéntica al documento de referencia (Matriz_Casos_Prueba_SafeIndustrial.docx)
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Márgenes ────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'),  'single')
        tag.set(qn('w:sz'),   '4')
        tag.set(qn('w:space'),'0')
        tag.set(qn('w:color'),'AAAAAA')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def cell_txt(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(text, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

# ── Colores ───────────────────────────────────────────────────────────────────
COL_HEADER  = '2E74B5'
COL_AREA    = 'D6E4F7'
COL_AVANZA  = 'D5E8D4'
COL_FALLA   = 'FFD7D7'
COL_BLOQ    = 'FFF3CD'
COL_WHITE   = 'FFFFFF'
COL_GRAY    = 'F2F2F2'

# ═════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Matriz de Casos de Prueba')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Proyecto: Plataforma SafeIndustrial — Backend Integrador\nMinera / Industria AuriMax S.A.C.')
run2.font.size = Pt(13)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(
    'Integrantes: Jefferson Calderón Ponce · Ricardo Palomino Contreras · Imanol Quispe Aguilar\n'
    'Ciclo 09 — Arequipa, Perú — 2026'
)
run3.font.size = Pt(10)
run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  CRITERIOS DE VALIDACIÓN
# ═════════════════════════════════════════════════════════════════════════════
heading('Criterios de Validación de Estado', level=2)

crit = doc.add_table(rows=4, cols=2)
crit.style = 'Table Grid'
crit.alignment = WD_TABLE_ALIGNMENT.LEFT

hdr_row = crit.rows[0]
set_cell_bg(hdr_row.cells[0], COL_HEADER)
set_cell_bg(hdr_row.cells[1], COL_HEADER)
cell_txt(hdr_row.cells[0], 'Estado', bold=True, color='FFFFFF', size=10)
cell_txt(hdr_row.cells[1], 'Descripción', bold=True, color='FFFFFF', size=10)

estados = [
    ('Avanza',   COL_AVANZA, 'La prueba existe, se ejecuta y el resultado obtenido coincide con el esperado.'),
    ('Falla',    COL_FALLA,  'La prueba no existe o el módulo no está implementado (sin cobertura de prueba).'),
    ('Bloqueado',COL_BLOQ,  'No se puede probar porque el módulo o infraestructura aún no está implementada.'),
]
for i, (e, c, d) in enumerate(estados, 1):
    set_cell_bg(crit.rows[i].cells[0], c)
    set_cell_bg(crit.rows[i].cells[1], COL_WHITE)
    cell_txt(crit.rows[i].cells[0], e, bold=True, size=10)
    cell_txt(crit.rows[i].cells[1], d, size=10)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 1 — MATRIZ GLOBAL
# ═════════════════════════════════════════════════════════════════════════════
heading('1. Matriz Global de Casos de Prueba (71 TCs)', level=1)
body('La siguiente tabla consolida todos los casos de prueba del proyecto, organizados por área empresarial. '
     'Los casos con estado Falla o Bloqueado indican requisitos sin cobertura de prueba en el backend actual.', size=10)
doc.add_paragraph()

# Cabecera de la tabla global
COLS_GLOBAL = [
    'ID', 'Requerimiento', 'Tipo de Prueba', 'Escenario',
    'Datos de Entrada', 'Resultado Esperado', 'Resultado Obtenido', 'Estado'
]
COL_WIDTHS = [Cm(1.5), Cm(3.2), Cm(2.4), Cm(4.0), Cm(3.2), Cm(3.8), Cm(3.8), Cm(1.8)]

def make_global_table():
    tbl = doc.add_table(rows=1, cols=len(COLS_GLOBAL))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Widths
    for i, w in enumerate(COL_WIDTHS):
        tbl.columns[i].width = w
    # Header row
    hr = tbl.rows[0]
    for i, h in enumerate(COLS_GLOBAL):
        set_cell_bg(hr.cells[i], COL_HEADER)
        cell_txt(hr.cells[i], h, bold=True, color='FFFFFF', size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    return tbl

def add_area_row(tbl, area_label):
    row = tbl.add_row()
    for i in range(len(COLS_GLOBAL)):
        set_cell_bg(row.cells[i], COL_AREA)
        set_cell_border(row.cells[i])
    # Merge all cells and write area name
    row.cells[0].merge(row.cells[-1])
    cell_txt(row.cells[0], area_label, bold=True, size=9, color='1F497D',
             align=WD_ALIGN_PARAGRAPH.LEFT)

def estado_color(e):
    if e == 'Avanza':    return COL_AVANZA
    if e == 'Falla':     return COL_FALLA
    if e == 'Bloqueado': return COL_BLOQ
    return COL_WHITE

def add_tc_row(tbl, tc_id, req, tipo, escenario, datos, esperado, obtenido, estado):
    row = tbl.add_row()
    vals = [tc_id, req, tipo, escenario, datos, esperado, obtenido, estado]
    bg = estado_color(estado)
    for i, v in enumerate(vals):
        set_cell_bg(row.cells[i], COL_GRAY if i % 2 == 0 else COL_WHITE)
        set_cell_border(row.cells[i])
    # Estado column always colored
    set_cell_bg(row.cells[-1], bg)
    for i, v in enumerate(vals):
        cell_txt(row.cells[i], v, size=8)

# ── DATOS DE LA TABLA GLOBAL ──────────────────────────────────────────────────
tcs = {
    'A': [
        # ID, Requerimiento, Tipo, Escenario, Datos, Esperado, Obtenido, Estado
        ('TC-A01','RF001 – Catálogo de Cursos','Funcional','Listar cursos disponibles',
         'Sin auth requerida','Lista de cursos con título, precio, instructor','Cursos devueltos correctamente','Avanza'),
        ('TC-A02','RF001 – Catálogo de Cursos','Funcional','Curso inexistente por ID',
         'ID inválido','Error 404 / ResourceNotFoundException','Excepción lanzada correctamente','Avanza'),
        ('TC-A03','RF002 – Gestión Promociones','Funcional','Crear cupón nuevo',
         'Código "SAVE10", % descuento','Cupón creado con código en mayúsculas','Cupón persistido exitosamente','Avanza'),
        ('TC-A04','RF002 – Gestión Promociones','Funcional','Cupón con código duplicado',
         'Código ya existente','Error de validación / DataIntegrityViolation','Excepción lanzada','Avanza'),
        ('TC-A05','RF002 – Gestión Promociones','Funcional','Cupón expirado',
         'Fecha de vencimiento pasada','Error "cupón expirado"','CouponExpiredException lanzada','Avanza'),
        ('TC-A06','RF002 – Gestión Promociones','Funcional','Toggle estado cupón activo→inactivo',
         'Cupón ACTIVE','Estado cambia a DISABLED','Estado DISABLED confirmado','Avanza'),
        ('TC-A07','RF003 – Carrito de Compras','Funcional','Crear orden sin cupón',
         'Ítems válidos + total > 0','Orden creada con estado PENDING','Orden persistida, evento publicado','Avanza'),
        ('TC-A08','RF003 – Carrito de Compras','Funcional','Crear orden con cupón PERCENTAGE',
         'Ítems + cupón SAVE10','Descuento aplicado, total reducido','Total con deducción calculado','Avanza'),
        ('TC-A09','RF003 – Carrito de Compras','Funcional','Orden con lista de ítems vacía',
         'items=[]','Error de validación','IllegalArgumentException lanzada','Avanza'),
        ('TC-A10','RF003 – Carrito de Compras','Funcional','Cancelar orden en estado PENDING',
         'Orden PENDING','Orden cancelada','Estado CANCELLED guardado','Avanza'),
        ('TC-A11','RF004 – Comprobante de Pago','Funcional','Pago aprobado por MercadoPago',
         'Token MP válido, monto > 0','Pago completado, estado SUCCESS','Evento emitido correctamente','Avanza'),
        ('TC-A12','RF004 – Comprobante de Pago','Funcional','Pago rechazado por MercadoPago',
         'Token MP rechazado','Pago fallido, estado FAILED','Estado FAILED registrado','Avanza'),
        ('TC-A13','RF004 – Notificación email compra','Funcional','Email de confirmación de compra exitosa',
         'Datos de orden + email alumno','Email enviado al comprador','JavaMailSender invocado correctamente','Avanza'),
    ],
    'B': [
        ('TC-B01','RF005 – Solicitud Compra EPP','Funcional','Crear solicitud de compra válida',
         'categoria, cantidad >= 1','Solicitud creada con estado PENDIENTE','Estado PENDIENTE persistido','Avanza'),
        ('TC-B02','RF005 – Solicitud Compra EPP','Funcional','Obtener solicitud por ID',
         'ID existente','PurchaseRequestResponse devuelto','Respuesta mapeada correctamente','Avanza'),
        ('TC-B03','RF005 – Solicitud Compra EPP','Funcional','Solicitud no encontrada',
         'ID inválido','Error 404 / ResourceNotFoundException','Excepción lanzada','Avanza'),
        ('TC-B04','RF005 – Solicitud Compra EPP','Funcional','Actualizar estado a APROBADO',
         'ID + estado: "APROBADO"','Estado cambiado a APROBADO','Estado actualizado y persistido','Avanza'),
        ('TC-B05','RF005 – Solicitud Compra EPP','Funcional','Estadísticas del sistema de compras',
         'Solicitudes en BD','Totales por estado calculados','StatsResponse correcto','Avanza'),
        ('TC-B06','RF005 – Solicitud Compra EPP','Funcional','Listar solicitudes aprobadas',
         'Mix de estados en BD','Solo solicitudes APROBADO','Lista filtrada correctamente','Avanza'),
        ('TC-B07','RF006 – Inventario de EPP','Funcional','Listar todos los ítems de inventario',
         'Sin filtro','Lista de InventoryItem','Lista devuelta correctamente','Avanza'),
        ('TC-B08','RF006 – Inventario de EPP','Funcional','Crear ítem de inventario',
         'codigo, descripcion, stock >= 0','Ítem creado y persistido','InventoryItemResponse devuelto','Avanza'),
        ('TC-B09','RF006 – Inventario de EPP','Funcional','Crear ítem con codigo vacío',
         'codigo: ""','Error 400 de validación','MethodArgumentNotValidException','Avanza'),
        ('TC-B10','RF007 – Entrega EPP a trabajador','Funcional','Entregar EPP con stock suficiente',
         'inventoryItemId + cantidad <= stock','Stock decrementado, entrega registrada','EppDelivery persistida, evento publicado','Avanza'),
        ('TC-B11','RF007 – Entrega EPP a trabajador','Funcional','Stock insuficiente para entrega',
         'cantidad > stock disponible','Error 409 / InsufficientStockException','Excepción lanzada','Avanza'),
        ('TC-B12','RF007 – Entrega EPP a trabajador','Funcional','Solicitud de compra no encontrada',
         'inventoryItemId inválido','Error 404','ResourceNotFoundException lanzada','Avanza'),
        ('TC-B13','RFNN – EPP asignado por DNI','Funcional','Ver entregas del trabajador por DNI',
         'workerDni existente','Lista de EppDeliveryResponse','Entregas filtradas por DNI','Avanza'),
        ('TC-B14','RFNN – EPP asignado por DNI','Funcional','Trabajador sin entregas devuelve vacío',
         'workerDni sin datos','Lista vacía []','Lista vacía retornada','Avanza'),
    ],
    'C': [
        ('TC-C01','RF008 – Gestión Cursos','Funcional','Crear curso con secciones y lecturas',
         'Datos curso + secciones','Curso persistido con UUIDs únicos','Curso creado, caché actualizado','Avanza'),
        ('TC-C02','RF008 – Gestión Cursos','Funcional','Obtener curso por ID válido',
         'ID de curso existente','CourseResponse con datos completos','Curso mapeado correctamente','Avanza'),
        ('TC-C03','RF008 – Gestión Cursos','Funcional','Actualizar curso existente',
         'Datos modificados + ID','Curso actualizado, caché refrescado','Guardado con nuevos valores','Avanza'),
        ('TC-C04','RF008 – Gestión Cursos','Funcional','Eliminar curso por ID',
         'ID curso válido','Curso eliminado, caché limpiado','Delete ejecutado, evict confirmado','Avanza'),
        ('TC-C05','RF009 – Evaluaciones','Funcional','Crear examen con preguntas',
         'Lista de preguntas','Examen creado con orderIndex asignado','Examen persistido con índices','Avanza'),
        ('TC-C06','RF009 – Evaluaciones','Funcional','Enviar intento con respuestas',
         'Respuestas del alumno','Intento calificado automáticamente','Resultado calculado y guardado','Avanza'),
        ('TC-C07','RF010 – Certificación Digital','Funcional','Alumno aprueba examen (nota >= mínimo)',
         'Respuestas correctas','Certificado PDF generado y enviado','Cert generado + email enviado','Avanza'),
        ('TC-C08','RF010 – Certificación Digital','Funcional','Alumno reprueba examen (nota < mínimo)',
         'Respuestas incorrectas','Sin certificado generado','No se crea cert, no se envía email','Avanza'),
        ('TC-C09','RF011 – Registro Incidentes IA','Funcional','Crear incidente EPP desde cámara',
         'cameraKey, violationTypes, confidence','Incidente creado con estado PENDING','Incidente persistido','Avanza'),
        ('TC-C10','RF011 – Registro Incidentes IA','Funcional','Crear incidente con campos faltantes',
         'Payload incompleto','Error 400 de validación','MethodArgumentNotValidException','Avanza'),
        ('TC-C11','RF012 – Revisión Incidentes','Funcional','Jefe aprueba incidente con workerId',
         'status=APPROVED + workerId','Puntos descontados, alerta publicada','ComplianceScore actualizado','Avanza'),
        ('TC-C12','RF012 – Revisión Incidentes','Funcional','Jefe rechaza incidente (falso positivo)',
         'status=REJECTED','Sin deducción ni alerta','Estado REJECTED guardado','Avanza'),
        ('TC-C13','RF012 – Revisión Incidentes','Funcional','Aprobar sin workerId lanza error',
         'status=APPROVED + workerId=null','Error 400','IllegalArgumentException','Avanza'),
        ('TC-C14','RF013 – Puntaje Cumplimiento','Funcional','Trabajador consulta su puntaje',
         'X-User-Id registrado','WorkerComplianceScoreResponse','Score devuelto','Avanza'),
        ('TC-C15','RF013 – Puntaje Cumplimiento','Funcional','Trabajador nuevo recibe puntaje base',
         'X-User-Id sin registro','Score = 100 (base)','Score base retornado','Avanza'),
        ('TC-C16','RF014 – Apelaciones','Funcional','Trabajador apela infracción propia',
         'incidente APPROVED + mismo workerId','appealStatus = PENDING','Apelación registrada','Avanza'),
        ('TC-C17','RF014 – Apelaciones','Funcional','Jefe aprueba apelación — restaura puntos',
         'approved=true','Score restaurado, estado APPEALED','ComplianceScore actualizado','Avanza'),
        ('TC-C18','RF014 – Apelaciones','Funcional','Jefe rechaza apelación — infracción vigente',
         'approved=false','Estado REJECTED, score sin cambio','Estado REJECTED guardado','Avanza'),
        ('TC-C19','RFNN – Monitoreo Cámaras IA','Funcional','Detectar infracción EPP con YOLOv11',
         'Frame de cámara IP','Alerta generada con foto evidencia','Sin prueba implementada','Bloqueado'),
        ('TC-C20','RFNN – Notif. Tiempo Real','Funcional','Alerta WebSocket al supervisor',
         'Evento de detección IA','Notificación enviada a topic WS','SimpMessagingTemplate invocado','Avanza'),
        ('TC-C21','RFNN – Ranking Seguridad','No Funcional','Calcular ranking de cumplimiento EPP',
         'Datos de infracciones por trabajador','Ranking ordenado por área','Sin prueba implementada','Bloqueado'),
    ],
    'D': [
        ('TC-D01','RF015 – Control Acceso RBAC','Funcional','Crear usuario con rol asignado',
         'Email, nombre, rol "supervisor"','Usuario creado en BD + Keycloak','Usuario guardado, Keycloak invocado','Avanza'),
        ('TC-D02','RF015 – Control Acceso RBAC','Funcional','Activar/desactivar cuenta de usuario',
         'ID usuario activo','Estado cambia a INACTIVE','Toggle ejecutado, Keycloak sincronizado','Avanza'),
        ('TC-D03','RF015 – Control Acceso RBAC','Seguridad','Registro con email duplicado',
         'Email ya registrado','Error de validación','DataIntegrityViolationException','Avanza'),
        ('TC-D04','RF015 – Control Acceso RBAC','Funcional','Sincronización Keycloak ID al login',
         'keycloakId diferente en BD','ID sincronizado automáticamente','Usuario actualizado','Avanza'),
        ('TC-D05','RF016 – Foro y Mensajería','Funcional','Crear conversación entre usuarios',
         'ConversationRequest válido','Conversación creada y persistida','ConversationResponse devuelto','Avanza'),
        ('TC-D06','RF016 – Foro y Mensajería','Funcional','Enviar mensaje en conversación',
         'MessageRequest + conversationId','Mensaje guardado','MessageResponse devuelto','Avanza'),
        ('TC-D07','RF016 – Foro y Mensajería','Funcional','Crear post de foro',
         'ForumPostRequest válido','Post creado y persistido','ForumPostResponse devuelto','Avanza'),
        ('TC-D08','RF016 – Foro y Mensajería','Funcional','Listar posts de un curso',
         'courseId existente','Lista de ForumPostResponse','Posts filtrados por curso','Avanza'),
        ('TC-D09','RF017 – Notificaciones','Funcional','Enviar email de confirmación',
         'Datos de orden + email alumno','Email enviado vía JavaMailSender','JavaMailSender invocado','Avanza'),
        ('TC-D10','RF017 – Notificaciones','Funcional','Publicar alerta WebSocket',
         'WebAlertEvent válido','Alerta enviada por WS','SimpMessagingTemplate invocado','Avanza'),
        ('TC-D11','RF017 – Notificaciones','Funcional','Health check del servicio notificaciones',
         'GET /actuator/health','200 OK con estado UP','Estado UP confirmado','Avanza'),
        ('TC-D12','RFNN – Dashboard KPIs','No Funcional','Visualizar KPIs de seguridad',
         'Usuario gerente autenticado','Dashboard con indicadores consolidados','Sin prueba implementada','Bloqueado'),
        ('TC-D13','RFNN – Auditoría y Log','No Funcional','Registrar evento de auditoría',
         'Acción + timestamp','Log registrado con trazabilidad','Sin prueba implementada','Bloqueado'),
    ]
}

area_labels = {
    'A': 'Área A — Marketing y Ventas',
    'B': 'Área B — Logística y Almacén',
    'C': 'Área C — Producción y Operaciones',
    'D': 'Área D — Dirección o Gerencia General',
}

tbl = make_global_table()
for area, rows in tcs.items():
    add_area_row(tbl, area_labels[area])
    for r in rows:
        add_tc_row(tbl, *r)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 2 — DETALLE DE REQUISITOS POR ÁREA
# ═════════════════════════════════════════════════════════════════════════════
heading('2. Detalle de Requisitos por Área Empresarial', level=1)

areas_req = {
    '2.1 Área A — Marketing y Ventas': {
        'desc': (
            'Comprende los requisitos relacionados con la publicación de cursos, '
            'gestión de promociones y descuentos, proceso de compra (órdenes) y '
            'generación de comprobantes de pago via MercadoPago.'
        ),
        'nota': None,
        'reqs': [
            ('RF001','Catálogo público de cursos',
             'GET /api/v1/courses\nGET /api/v1/courses/{id}','Público / Alumno'),
            ('RF002','Gestión de promociones y descuentos (cupones)',
             'POST /api/v1/coupons\nGET /api/v1/coupons\nDELETE /api/v1/coupons/{id}\nPATCH /api/v1/coupons/{id}/toggle',
             'Gerente / Marketing'),
            ('RF003','Carrito de compras y órdenes',
             'POST /api/v1/orders\nGET /api/v1/orders/{id}\nDELETE /api/v1/orders/{id}','Alumno'),
            ('RF004','Generación de comprobantes de pago',
             'POST /api/v1/payments/process (automático)','Sistema (automático)'),
        ]
    },
    '2.2 Área B — Logística y Almacén': {
        'desc': (
            'Comprende los requisitos de gestión de inventario de EPP (Equipos de Protección Personal), '
            'solicitudes de compra, entregas a trabajadores y visualización de EPP asignado.'
        ),
        'nota': None,
        'reqs': [
            ('RF005','Solicitud de compra de EPP',
             'POST /api/v1/purchase/requests\nGET /api/v1/purchase/requests\nGET /api/v1/purchase/requests/{id}\nPUT /api/v1/purchase/requests/{id}\nGET /api/v1/purchase/requests/stats\nGET /api/v1/purchase/requests/inventory',
             'Almacenero / Gerente'),
            ('RF006','Inventario de EPP',
             'GET /api/v1/purchase/inventory\nPOST /api/v1/purchase/inventory','Almacenero'),
            ('RF007','Asignación y entrega de EPP a trabajadores',
             'POST /api/v1/purchase/epp/deliver\nGET /api/v1/purchase/epp/worker','Almacenero'),
            ('RFNN','Visualización de EPP asignado al trabajador',
             'GET /api/v1/purchase/epp/deliveries?workerDni={dni}','Trabajador / Supervisor'),
        ]
    },
    '2.3 Área C — Producción y Operaciones': {
        'desc': (
            'Comprende los requisitos de gestión de cursos por instructores, evaluaciones, '
            'certificación digital automática, registro de incidentes EPP desde cámaras IA (YOLOv11), '
            'revisión de infracciones, puntaje de cumplimiento y sistema de apelaciones.'
        ),
        'nota': None,
        'reqs': [
            ('RF008','Gestión de cursos por instructor',
             'GET /api/v1/courses\nPOST /api/v1/courses\nPUT /api/v1/courses/{id}\nDELETE /api/v1/courses/{id}',
             'Instructor / Gerente'),
            ('RF009','Gestión de evaluaciones y analítica',
             'POST /api/v1/exams\nGET /api/v1/exams/{id}\nPOST /api/v1/exams/{id}/attempt','Instructor / Alumno'),
            ('RF010','Certificación digital automática',
             'POST /api/v1/exams/{id}/attempt/submit (automático)','Sistema (automático)'),
            ('RF011','Registro de incidentes EPP por cámara IA',
             'POST /api/v1/incidents\nGET /api/v1/incidents\nGET /api/v1/incidents/{id}\nGET /api/v1/incidents/mine',
             'Sistema IA / Trabajador / Supervisor'),
            ('RF012','Revisión y aprobación de incidentes',
             'PATCH /api/v1/incidents/{id}/review','Jefe de Seguridad / Supervisor'),
            ('RF013','Puntaje de cumplimiento EPP del trabajador',
             'GET /api/v1/safety-score/me','Trabajador'),
            ('RF014','Sistema de apelaciones de infracciones',
             'POST /api/v1/incidents/{id}/appeal\nGET /api/v1/incidents/appeals\nPATCH /api/v1/incidents/{id}/appeal/resolve',
             'Trabajador / Jefe de Seguridad'),
            ('RFNN','Monitoreo de cámaras IA (YOLOv11 — servicio Python)',
             'WebSocket / ws/alerts\nGET /api/v1/safety/alerts','Supervisor (Bloqueado)'),
            ('RFNN','Panel de alertas en tiempo real',
             'WebSocket topic: /topic/alerts\nSimpMessagingTemplate','Supervisor'),
            ('RFNN','Ranking de seguridad / Gamificación',
             'GET /api/v1/safety/ranking (no implementado)','Supervisor / Gerente (Bloqueado)'),
        ]
    },
    '2.4 Área D — Dirección o Gerencia General': {
        'desc': (
            'Comprende los requisitos de control de acceso basado en roles (RBAC con Keycloak), '
            'foro y mensajería interna, notificaciones email y WebSocket, dashboard de KPIs '
            'y auditoría de logs del sistema.'
        ),
        'nota': None,
        'reqs': [
            ('RF015','Control de acceso y gestión de usuarios (RBAC)',
             'GET /api/v1/users\nPOST /api/v1/users\nPATCH /api/v1/users/{id}\nDELETE /api/v1/users/{id}\n+ Integración Keycloak',
             'Gerente / Sistema'),
            ('RF016','Foro y mensajería interna entre usuarios',
             'POST /api/v1/conversations\nGET /api/v1/conversations/{id}\nPOST /api/v1/conversations/{id}/messages\nPOST /api/v1/forum/posts\nGET /api/v1/forum/posts',
             'Alumno / Instructor / Supervisor'),
            ('RF017','Notificaciones email y WebSocket',
             'Consumidor RabbitMQ (automático)\nGET /actuator/health','Sistema (automático)'),
            ('RF010-M','Módulo de aprobaciones de precio de cursos',
             'PATCH /api/v1/courses/{id}/price-request\nPATCH /api/v1/courses/{id}/price-approve','Gerente'),
            ('RFNN','Dashboard general de KPIs de seguridad',
             'Frontend Redux (sin endpoint backend directo)','Gerente (Bloqueado)'),
            ('RFNN','Auditoría y Log del sistema',
             'GET /api/v1/audit/logs (no implementado)','Gerente (Bloqueado)'),
        ]
    },
}

for area_title, area_data in areas_req.items():
    subheading(area_title)
    body(area_data['desc'], size=10)
    if area_data['nota']:
        p = doc.add_paragraph()
        run = p.add_run('NOTA: ' + area_data['nota'])
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    req_tbl = doc.add_table(rows=1, cols=4)
    req_tbl.style = 'Table Grid'
    req_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    req_tbl.columns[0].width = Cm(1.8)
    req_tbl.columns[1].width = Cm(6.5)
    req_tbl.columns[2].width = Cm(8.5)
    req_tbl.columns[3].width = Cm(3.0)

    hdr = req_tbl.rows[0]
    for i, h in enumerate(['ID', 'Requisito Funcional', 'Endpoint / Origen', 'Rol Autorizado']):
        set_cell_bg(hdr.cells[i], COL_HEADER)
        cell_txt(hdr.cells[i], h, bold=True, color='FFFFFF', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for req in area_data['reqs']:
        rrow = req_tbl.add_row()
        for i, v in enumerate(req):
            bg = COL_AREA if i == 0 else (COL_GRAY if i % 2 == 0 else COL_WHITE)
            set_cell_bg(rrow.cells[i], bg)
            set_cell_border(rrow.cells[i])
            cell_txt(rrow.cells[i], v, bold=(i == 0), size=9)

    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 3 — REQUISITOS SIN COBERTURA
# ═════════════════════════════════════════════════════════════════════════════
heading('3. Requisitos sin Cobertura de Prueba', level=1)
body(
    'Los siguientes requisitos funcionales NO cuentan con pruebas unitarias ni de integración '
    'implementadas en el backend actual. Se recomienda agregar pruebas o implementar los módulos faltantes.',
    size=10
)
doc.add_paragraph()

sin_cobertura = [
    ('RFNN – Monitoreo Cámaras IA (YOLOv11)',
     'El servicio Python de detección visual no tiene microservicio Spring con pruebas. '
     'El ingreso de incidentes es manual o vía API POST /incidents. '
     'La integración directa con la cámara no puede automatizarse en tests.'),
    ('RFNN – Ranking de Seguridad / Gamificación',
     'Sin implementación de endpoint GET /api/v1/safety/ranking ni lógica de servicio. '
     'No existe prueba unitaria ni de integración.'),
    ('RFNN – Reporte de Cuasi Accidentes',
     'Sin endpoint POST /api/v1/safety/near-miss implementado. '
     'No existe entidad, repositorio ni prueba.'),
    ('RF011 – Dashboard General de KPIs',
     'El dashboard de gerencia no tiene endpoint backend propio. '
     'Los datos se obtienen de múltiples servicios vía Redux en el frontend. '
     'Sin prueba de controlador implementada.'),
    ('RFNN – Auditoría y Log del Sistema',
     'El endpoint GET /api/v1/audit/logs no está implementado en ningún microservicio. '
     'No existe prueba unitaria ni de integración.'),
    ('RFNN – Alertas WebSocket (prueba de integración real)',
     'Las pruebas de WebSocket con SimpMessagingTemplate están cubiertos con mocks unitarios '
     'pero no con pruebas de integración WebSocket real (falta @SpringBootTest con puerto real + StompClient).'),
]

for titulo, detalle in sin_cobertura:
    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run(titulo + ': ')
    run1.bold = True
    run1.font.size = Pt(10)
    run2 = p.add_run(detalle)
    run2.font.size = Pt(10)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 4 — RESUMEN DE COBERTURA POR MICROSERVICIO
# ═════════════════════════════════════════════════════════════════════════════
heading('4. Resumen de Cobertura por Microservicio', level=1)

cov_data = [
    ('safety-service',       '43', '34', '≥ 80%', 'Implementado'),
    ('purchase-service',     '28', '22', '≥ 80%', 'Implementado'),
    ('order-service',        '16', '10', '≥ 80%', 'Implementado'),
    ('course-service',       '6',  '5',  '≥ 80%', 'Implementado'),
    ('exam-service',         '8',  '4',  '≥ 80%', 'Implementado'),
    ('user-service',         '7',  '5',  '≥ 80%', 'Implementado'),
    ('payment-service',      '4',  '4',  '≥ 80%', 'Implementado'),
    ('chat-service',         '8',  '6',  '≥ 80%', 'Implementado'),
    ('notification-service', '4',  '2',  '≥ 80%', 'Implementado'),
    ('TOTAL',               '124', '92',  '—',    '—'),
]

cov_tbl = doc.add_table(rows=1, cols=5)
cov_tbl.style = 'Table Grid'
cov_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, h in enumerate(['Microservicio','Tests Unitarios','Tests Integración','Cobertura Objetivo','Estado']):
    set_cell_bg(cov_tbl.rows[0].cells[i], COL_HEADER)
    cell_txt(cov_tbl.rows[0].cells[i], h, bold=True, color='FFFFFF', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

for row_data in cov_data:
    row = cov_tbl.add_row()
    for i, v in enumerate(row_data):
        bg = COL_GRAY if i == 0 else COL_WHITE
        if row_data[0] == 'TOTAL': bg = COL_AREA
        set_cell_bg(row.cells[i], bg)
        set_cell_border(row.cells[i])
        cell_txt(row.cells[i], v, bold=(row_data[0]=='TOTAL'), size=9,
                 align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()

# ── Pie de página ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Fin del documento —')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Guardar ───────────────────────────────────────────────────────────────────
OUT = r'C:\Users\panc1\Downloads\Matriz_Casos_Prueba_SafeIndustrial_Completa.docx'
doc.save(OUT)
print(f'Documento generado: {OUT}')
